# notice_app.py (single file, refactored)

import os, json, datetime
import tkinter as tk
import sys

def _app_base_dir():
    # PyInstaller 실행 파일 내부/외부 모두 대응
    if getattr(sys, 'frozen', False) and hasattr(sys, '_MEIPASS'):
        # onefile 모드에서 임시 해제 폴더
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))

APP_DIR   = _app_base_dir()
DATA_DIR  = os.path.join(os.getenv("LOCALAPPDATA", APP_DIR), "ChildcareNoticeMaker")
OUT_DIR   = os.path.join(DATA_DIR, "output")        # 기존 OUT_DIR 대체
STATE_FILE = os.path.join(DATA_DIR, "rotation_state.json")  # 기존 STATE_FILE 대체

os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(OUT_DIR,  exist_ok=True)

from tkinter import ttk, messagebox, filedialog

APP_NAME   = "Childcare Notice Maker"
OUT_DIR    = "output"
STATE_FILE = "rotation_state.json"

# -------------------- python-docx (optional) --------------------
try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except Exception:
    # 설치 안내: pip install python-docx
    DOCX_AVAILABLE = False

# -------------------- 스타일/페이지 설정 상수 --------------------
BASE_FONT          = "Malgun Gothic"  # Windows 한글 기본: 맑은 고딕
BASE_SIZE          = 11               # 본문 pt
TITLE_SIZE         = 28               # 제목 pt
SECTION_GAP_BEFORE = 6                # '■' 섹션 위쪽 여백(pt)
SECTION_GAP_AFTER  = 2                # '■' 섹션 아래 여백(pt)
PARA_GAP_AFTER     = 2                # 일반 문단 아래 여백(pt)

# -------------------- 시즌/로테이션 문구 --------------------
SEASON_LINES = {
    "spring": [
        "포근한 봄바람 속에서 아이들의 마음도 살포시 피어나는 계절입니다.",
        "꽃내음 가득한 길을 걸으며 작은 발견 하나하나가 배움이 되길 바랍니다.",
        "새로운 시작을 응원하며, 하루하루의 설렘을 조심스레 품어 보겠습니다."
    ],
    "summer": [
        "반짝이는 햇살처럼 아이들의 웃음도 환한 계절입니다. 건강과 안전을 먼저 챙기겠습니다.",
        "시원한 쉼과 알찬 놀이가 균형을 이루도록 일정과 환경을 세심히 준비했습니다.",
        "더운 날씨에도 마음은 가볍게, 물 한 모금과 그늘 한 자리까지 꼼꼼히 살피겠습니다."
    ],
    "autumn": [
        "높고 맑은 하늘 아래 아이들의 하루가 고운 빛으로 물들어 갑니다.",
        "바스락거리는 낙엽처럼 작은 변화도 소중히 담아 따뜻한 배움으로 이어가겠습니다.",
        "수확의 기쁨처럼 노력의 열매가 맺어지도록 차분히 동행하겠습니다."
    ],
    "winter": [
        "차가운 바람 속에서도 교실은 늘 포근하도록, 안전과 건강을 먼저 살피겠습니다.",
        "따뜻한 마음과 정돈된 일과로 아이들의 하루를 든든하게 채우겠습니다.",
        "서늘한 계절에도 작은 성취를 꼼꼼히 기록하며 잔잔한 기쁨을 나누겠습니다."
    ]
}

INTRO_VARIANTS = {
    "Field Trip": [
        "아이들이 직접 보고 듣고 만지며 배우는 시간이 될 수 있도록 차분히 준비했습니다. 안전을 가장 먼저 생각하며, 아이 한 명 한 명의 속도에 맞춰 살피겠습니다.",
        "자연과 사회를 가까이에서 경험하는 하루입니다. 친구와 배려하고 질서를 지키는 작은 습관까지 함께 익힐 수 있도록 교사들이 곁에서 세심히 도우겠습니다.",
        "설렘 가득한 마음이 배움으로 이어지도록 일정과 동선을 꼼꼼히 구성했습니다. 무사히 다녀올 수 있도록 가정의 격려와 협조를 부탁드립니다."
    ],
    "Class Observation": [
        "평소 교실에서의 배움과 놀이가 어떻게 흐르는지, 아이들이 어떤 표정으로 하루를 보내는지 천천히 살펴보실 수 있도록 준비했습니다.",
        "가정과 원이 같은 마음으로 아이를 바라볼 때 성장의 결이 고와집니다. 편안한 마음으로 오셔서 작은 변화도 함께 기뻐해 주세요.",
        "교사와 친구들 사이에서 나누는 따뜻한 상호작용을 가까이에서 보시며 가정과의 연계에 도움이 되길 바랍니다."
    ],
    "Picnic": [
        "계절의 숨결을 온몸으로 느끼며 자연 속에서 마음껏 뛰놀 수 있도록 소풍을 마련했습니다. 물 한 모금, 모자 하나까지 정성껏 챙기겠습니다.",
        "친구와 함께 걷고 나누는 시간이 예절과 배려를 단단히 세웁니다. 안전한 이동과 충분한 휴식으로 편안한 하루가 되게 하겠습니다.",
        "작은 잎 하나, 바람 한 줄기에도 놀라고 배우는 때입니다. 아이들의 기쁨이 오래 기억으로 남도록 세심히 동행하겠습니다."
    ],
    "Health/Immunization": [
        "건강은 하루의 배움과 생활을 지탱하는 첫걸음입니다. 필요한 확인을 차분히 진행하고자 하오니 몇 가지 안내에 협조 부탁드립니다.",
        "알레르기·복용 약 등 중요한 정보는 교실에서 곧바로 돌봄에 반영됩니다. 아이에게 가장 안전한 환경이 될 수 있도록 함께 살펴 주세요.",
        "작은 증상도 소중히 듣고 살피겠습니다. 궁금하신 점은 언제든 편히 말씀해 주시면 가정과 상의하며 좋은 선택을 하겠습니다."
    ],
    "General Notice": [
        "늘 아이들의 하루를 믿고 응원해 주셔서 고맙습니다. 필요한 소식만 차분히 정리해 드리오니 천천히 확인 부탁드립니다.",
        "가정과 원이 한마음으로 아이를 바라볼 때 하루가 더 따뜻해집니다. 안내를 살펴보시고 의견 주시면 더 세심히 반영하겠습니다.",
        "작은 약속을 지키는 힘이 큰 성장을 만듭니다. 투명하게 소통하며 믿음을 이어가겠습니다."
    ]
}

# -------------------- 로테이션 상태 --------------------
def load_state():
    if os.path.exists(STATE_FILE):
        try:
            with open(STATE_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {}

def save_state(state):
    with open(STATE_FILE, "w", encoding="utf-8") as f:
        json.dump(state, f, ensure_ascii=False, indent=2)

def pick_rotating(key, candidates, state):
    if not candidates:
        return ""
    idx = state.get(key, 0)
    text = candidates[idx % len(candidates)]
    state[key] = (idx + 1) % len(candidates)
    return text

def season_of(d):
    m = d.month
    if m in (3,4,5): return "spring"
    if m in (6,7,8): return "summer"
    if m in (9,10,11): return "autumn"
    return "winter"

# -------------------- 문안 생성 --------------------
def build_notice(params):
    today = datetime.date.today().strftime("%Y-%m-%d")
    d = datetime.datetime.strptime(params["date"], "%Y-%m-%d").date()

    if params.get("start") and params.get("end"):
        t_range = f"{params['date']} {params['start']}–{params['end']}"
    elif params.get("start") or params.get("end"):
        t_range = f"{params['date']} {params.get('start') or params.get('end')}"
    else:
        t_range = params['date']

    state = load_state()
    season = season_of(d)
    season_line = pick_rotating(f"SEASON_{season}", SEASON_LINES[season], state)
    intro_line  = pick_rotating(params["event_type"], INTRO_VARIANTS.get(params["event_type"], INTRO_VARIANTS["General Notice"]), state)
    save_state(state)

    long_intro = (
        f"사랑하는 {params['center']} {params['classname']} 학부모님께,\n\n"
        "안녕하십니까. 늘 저희 교육 활동에 따뜻한 관심을 보내주시는 모든 가정께 깊이 감사드립니다.\n"
        f"{season_line}\n{intro_line}\n"
    )

    def opt(label, v):
        return f"{label}{v}\n" if v and str(v).strip() else ""
    def fee(v):
        if not v or str(v).strip() in ("0 KRW","0원","0"): return ""
        return f"• 참가비: {v}\n"

    et   = params["event_type"]
    head = f"[가정통신문] {params['event_name'] or et}"
    foot = (
        "\n※ 위 일정 및 내용은 원 사정과 안전 상황에 따라 변경될 수 있습니다.\n"
        f"{today}\n{params['center']} 원장 {params['contact_name']} 드림"
    )

    if et == "Class Observation":
        body = (
            f"{head}\n\n{long_intro}\n"
            "■ 참관 수업 개요\n"
            f"• 일시: {t_range}\n"
            f"• 장소: {params['location']}\n"
            f"{opt('• 준비물: ', params['materials'])}"
            "\n■ 안내 및 협조 요청\n"
            "• 원내 안전을 위해 가정당 보호자 1인 참석을 권장드립니다.\n"
            "• 편안한 복장을 권하며, 외부 음식 반입은 자제 부탁드립니다.\n"
            "• 수업 중 임의 촬영은 제한되며, 별도 촬영본은 추후 공유드리겠습니다.\n"
            "\n■ 참석 회신\n"
            f"• 마감: {params['rsvp_deadline']}\n"
            "• 링크/회신 방법: (원에서 안내한 방식으로 회신)\n"
            f"{foot}"
        )
    elif et == "Field Trip":
        body = (
            f"{head}\n\n{long_intro}\n"
            "■ 체험학습 개요\n"
            f"• 일시: {t_range}\n"
            f"• 장소/프로그램: {params['location']}{opt(' / ', params['body_summary']).strip()}\n"
            f"{opt('• 이동수단: ', params['transport'])}{fee(params['cost'])}{opt('• 준비물/복장: ', params['materials'])}"
            "\n■ 안전 및 동의\n"
            "• 사진·영상 촬영 동의 및 안전수칙 확인이 필요합니다.\n"
            "• 알레르기/복용 약 등 특이사항은 반드시 기재해 주세요.\n"
            "\n■ 전자 동의서\n"
            f"• 마감: {params['rsvp_deadline']}\n"
            "• 링크/제출 방법: (원에서 안내한 방식으로 제출)\n"
            f"{foot}"
        )
    elif et == "Picnic":
        body = (
            f"{head}\n\n{long_intro}\n"
            "■ 소풍 개요\n"
            f"• 일시: {t_range}\n"
            f"• 장소: {params['location']}\n"
            f"{opt('• 준비물/복장: ', params['materials'])}{opt('• 우천 시 대체: ', params['rain_plan'])}{opt('• 이동수단: ', params['transport'])}"
            "\n■ 안내 및 협조 요청\n"
            "• 사진·영상 촬영 동의 및 안전수칙 확인을 부탁드립니다.\n"
            "• 알레르기/식단 제한이 있는 경우 간식 제공에 반영할 수 있도록 회신에 남겨 주세요.\n"
            "\n■ 전자 동의서\n"
            f"• 마감: {params['rsvp_deadline']}\n"
            "• 링크/제출 방법: (원에서 안내한 방식으로 제출)\n"
            f"{foot}"
        )
    elif et == "Health/Immunization":
        body = (
            f"{head}\n\n{long_intro}\n"
            "■ 진행 안내\n"
            f"• 일시/장소: {t_range} / {params['location']}\n"
            f"{opt('• 준비물: ', params['materials'])}"
            "\n■ 확인·동의 사항\n"
            "• 아이의 건강 상태, 알레르기, 복용 약 등을 정확히 기재해 주세요.\n"
            "• 필요한 경우 개별 상담을 진행하겠습니다.\n"
            "\n■ 확인서 제출\n"
            f"• 마감: {params['rsvp_deadline']}\n"
            "• 링크/제출 방법: (원에서 안내한 방식으로 제출)\n"
            f"{foot}"
        )
    else:
        body = (
            f"{head}\n\n{long_intro}\n"
            "■ 안내 내용\n"
            f"• {params['body_summary'] or params['materials'] or '상세 내용은 첨부를 참고해 주세요.'}\n"
            "\n■ 확인/회신(해당 시)\n"
            "• 링크/회신 방법: (원에서 안내한 방식으로 회신)\n"
            f"{foot}"
        )

    return body

# -------------------- 파일 저장 --------------------
def save_txt(content, path):
    os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)

# -------------------- DOCX 스타일러 --------------------
def _get_or_add_rPr(obj_element):
    """run._element / style._element에서 rPr 보장"""
    if hasattr(obj_element, "get_or_add_rPr"):
        return obj_element.get_or_add_rPr()
    rPr = getattr(obj_element, "rPr", None)
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        obj_element.insert(0, rPr)
    return rPr

def _get_or_add_rFonts(rPr):
    rFonts = getattr(rPr, "rFonts", None)
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    return rFonts

def _set_font_all_faces(obj_element, font_name):
    """ascii / hAnsi / eastAsia 모두 지정"""
    rPr    = _get_or_add_rPr(obj_element)
    rFonts = _get_or_add_rFonts(rPr)
    rFonts.set(qn("w:ascii"),   font_name)
    rFonts.set(qn("w:hAnsi"),   font_name)
    rFonts.set(qn("w:eastAsia"),font_name)

def add_rule(doc, width_chars=34):
    p = doc.add_paragraph("─" * width_chars)
    p.paragraph_format.space_after = Pt(PARA_GAP_AFTER)
    return p

def configure_page(section, margins_cm=(2,2,2,2)):
    section.orientation   = WD_ORIENT.PORTRAIT
    section.page_width    = Cm(21.0)
    section.page_height   = Cm(29.7)
    section.top_margin    = Cm(margins_cm[0])
    section.right_margin  = Cm(margins_cm[1])
    section.bottom_margin = Cm(margins_cm[2])
    section.left_margin   = Cm(margins_cm[3])

def add_header_image(doc, image_path):
    if not image_path or not os.path.exists(image_path):
        return
    sec = doc.sections[0]
    usable_width = sec.page_width - sec.left_margin - sec.right_margin
    doc.add_picture(image_path, width=usable_width)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

def build_styled_doc(content:str, output_path:str, header_image_path:str=None):
    """가정통신문 텍스트 → 스타일 적용 DOCX 생성"""
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx가 설치되어 있지 않습니다. 'pip install python-docx'를 먼저 실행해 주세요.")

    doc = Document()
    configure_page(doc.sections[0], margins_cm=(2,2,2,2))

    # Normal 스타일 전역
    style = doc.styles['Normal']
    style.font.name = BASE_FONT
    style.font.size = Pt(BASE_SIZE)
    _set_font_all_faces(style._element, BASE_FONT)

    # 불릿 스타일도 강제 (환경에 따라 없을 수 있음)
    try:
        lb = doc.styles['List Bullet']
        _set_font_all_faces(lb._element, BASE_FONT)
    except Exception:
        pass

    add_header_image(doc, header_image_path)

    lines = (content or "").splitlines()
    if not lines:
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        doc.save(output_path)
        return output_path

    # 제목(첫 비어있지 않은 줄)
    i = 0
    while i < len(lines) and not lines[i].strip():
        i += 1
    title_text = lines[i].strip() if i < len(lines) else "가정통신문"

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(title_text)
    run.font.name = BASE_FONT
    run.font.size = Pt(TITLE_SIZE)
    run.bold = True
    _set_font_all_faces(run._element, BASE_FONT)
    p_title.paragraph_format.space_after = Pt(6)

    add_rule(doc)

    # 본문
    for line in lines[i+1:]:
        t = (line or "").strip()

        if not t:
            p = doc.add_paragraph("")
            p.paragraph_format.space_after = Pt(PARA_GAP_AFTER)
            continue

        if t.startswith("■"):
            p = doc.add_paragraph(t)
            r = p.runs[0] if p.runs else p.add_run(t)
            r.font.name = BASE_FONT; r.font.size = Pt(BASE_SIZE+1); r.bold = True
            _set_font_all_faces(r._element, BASE_FONT)
            p.paragraph_format.space_before = Pt(SECTION_GAP_BEFORE)
            p.paragraph_format.space_after  = Pt(SECTION_GAP_AFTER)
            continue

        if t.startswith("•"):
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(t.lstrip("•").strip())
            r.font.name = BASE_FONT; r.font.size = Pt(BASE_SIZE)
            _set_font_all_faces(r._element, BASE_FONT)
            p.paragraph_format.space_after = Pt(0)
            continue

        if t.startswith("※"):
            p = doc.add_paragraph(t)
            r = p.runs[0] if p.runs else p.add_run(t)
            r.font.name = BASE_FONT; r.font.size = Pt(10)
            _set_font_all_faces(r._element, BASE_FONT)
            p.paragraph_format.space_after = Pt(2)
            continue

        if t.endswith("드림") or (t.count("-") == 2 and 8 <= len(t) <= 10):
            p = doc.add_paragraph(t)
            r = p.runs[0]
            r.font.name = BASE_FONT; r.font.size = Pt(11)
            _set_font_all_faces(r._element, BASE_FONT)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_after = Pt(PARA_GAP_AFTER)
            continue

        p = doc.add_paragraph(t)
        r = p.runs[0]
        r.font.name = BASE_FONT; r.font.size = Pt(BASE_SIZE)
        _set_font_all_faces(r._element, BASE_FONT)
        p.paragraph_format.space_after = Pt(PARA_GAP_AFTER)

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    return output_path

# -------------------- UI --------------------
class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title(APP_NAME)
        self.geometry("760x720")  # 버튼 가림 방지: 높이 ↑, 폭 약간 ↑
        self.resizable(False, False)

        # 필드
        self.center         = tk.StringVar(value="아이조아 어린이집")
        self.classname      = tk.StringVar(value="해바라기")
        self.contact_name   = tk.StringVar(value="원장 김OO")
        self.contact_phone  = tk.StringVar(value="010-1234-0000")

        self.event_type     = tk.StringVar(value="Field Trip")
        self.event_name     = tk.StringVar(value="체험학습 안내")
        self.date           = tk.StringVar(value=datetime.date.today().strftime("%Y-%m-%d"))
        self.start          = tk.StringVar(value="09:30")
        self.end            = tk.StringVar(value="14:00")
        self.location       = tk.StringVar(value="시립 자연학습원")
        self.materials      = tk.StringVar(value="물통, 모자")
        self.transport      = tk.StringVar(value="버스")
        self.cost           = tk.StringVar(value="0원")
        self.rsvp_deadline  = tk.StringVar(value=(datetime.date.today()+datetime.timedelta(days=5)).strftime("%Y-%m-%d 23:59"))
        self.rain_plan      = tk.StringVar(value="우천 시 실내 활동")
        self.body_summary   = tk.StringVar(value="세부 일정은 가정통신문 뒷면 참고")
        self.header_image_path = tk.StringVar(value="")  # 배너 이미지(선택)

        frm = ttk.Frame(self, padding=12)
        frm.pack(fill="both", expand=True)

        r=0
        def row(label, var, width=36):
            nonlocal r
            ttk.Label(frm, text=label, width=18, anchor="e").grid(row=r, column=0, padx=6, pady=4, sticky="e")
            e = ttk.Entry(frm, textvariable=var, width=width)
            e.grid(row=r, column=1, columnspan=3, sticky="w")
            r+=1

        ttk.Label(frm, text="행사유형", width=18, anchor="e").grid(row=r, column=0, padx=6, pady=4, sticky="e")
        cb = ttk.Combobox(frm, textvariable=self.event_type, values=list(INTRO_VARIANTS.keys()), width=34, state="readonly")
        cb.grid(row=r, column=1, sticky="w"); r+=1

        row("행사명/제목", self.event_name)
        row("어린이집 이름", self.center)
        row("반 이름", self.classname)
        row("일자 (YYYY-MM-DD)", self.date)
        row("시작시간 (HH:MM)", self.start)
        row("종료시간 (HH:MM)", self.end)
        row("장소", self.location)
        row("준비물", self.materials)
        row("이동수단", self.transport)
        row("참가비", self.cost)
        row("회신마감 (YYYY-MM-DD HH:MM)", self.rsvp_deadline)
        row("우천대응", self.rain_plan)
        row("요약(선택)", self.body_summary)
        row("원장 성함", self.contact_name)
        row("연락처", self.contact_phone)

        # 배너 이미지 선택
        img_frm = ttk.Frame(frm); img_frm.grid(row=r, column=0, columnspan=4, sticky="w", pady=(6,0))
        ttk.Label(img_frm, text="배너 이미지(선택):", width=18, anchor="e").grid(row=0, column=0, padx=6)
        ttk.Entry(img_frm, textvariable=self.header_image_path, width=42).grid(row=0, column=1, sticky="w")
        ttk.Button(img_frm, text="찾아보기", command=self._pick_image).grid(row=0, column=2, padx=6)
        r += 1

        # 버튼들
        btns = ttk.Frame(frm); btns.grid(row=r, column=0, columnspan=4, pady=12)
        ttk.Button(btns, text="문안 생성 (TXT)", command=self.make_txt).grid(row=0, column=0, padx=6)

        self.btn_docx = ttk.Button(btns, text="문안 + DOCX 생성", command=self.make_docx)
        self.btn_docx.grid(row=0, column=1, padx=6)
        if not DOCX_AVAILABLE:
            self.btn_docx.state(["disabled"])
            ttk.Label(frm, foreground="red",
                      text="※ python-docx 미설치: 'pip install python-docx' 후 DOCX 버튼 사용 가능").grid(row=r+1, column=0, columnspan=4, sticky="w")

        self.status = tk.StringVar(value="준비됨")
        ttk.Label(self, textvariable=self.status, relief="sunken", anchor="w").pack(fill="x", padx=12, pady=(0,8))

    def collect(self):
        return dict(
            center=self.center.get().strip(),
            classname=self.classname.get().strip(),
            contact_name=self.contact_name.get().strip(),
            contact_phone=self.contact_phone.get().strip(),
            event_type=self.event_type.get().strip(),
            event_name=self.event_name.get().strip(),
            date=self.date.get().strip(),
            start=self.start.get().strip(),
            end=self.end.get().strip(),
            location=self.location.get().strip(),
            materials=self.materials.get().strip(),
            transport=self.transport.get().strip(),
            cost=self.cost.get().strip(),
            rsvp_deadline=self.rsvp_deadline.get().strip(),
            rain_plan=self.rain_plan.get().strip(),
            body_summary=self.body_summary.get().strip()
        )

    def _save_paths(self, title_base):
        os.makedirs(OUT_DIR, exist_ok=True)
        safe = "".join(c for c in title_base if c.isalnum() or c in ("_","-"," ")).rstrip()
        txt_path  = os.path.join(OUT_DIR, f"{safe}.txt")
        docx_path = os.path.join(OUT_DIR, f"{safe}.docx")
        return txt_path, docx_path

    def make_txt(self):
        params = self.collect()
        try:
            content = build_notice(params)
            base = f"공지_{params['event_type']}__{params['classname']}_{params['date']}"
            txt_path, _ = self._save_paths(base)
            save_txt(content, txt_path)
            self.status.set(f"TXT 생성 완료 → {txt_path}")
            messagebox.showinfo(APP_NAME, f"문안(TXT) 생성 완료\n\n{txt_path}")
        except Exception as e:
            self.status.set("오류")
            messagebox.showerror(APP_NAME, str(e))

    def make_docx(self):
        params = self.collect()
        try:
            if not DOCX_AVAILABLE:
                raise RuntimeError("python-docx 미설치: 'pip install python-docx' 실행 후 다시 시도해 주세요.")
            content = build_notice(params)
            base = f"공지_{params['event_type']}__{params['classname']}_{params['date']}"
            txt_path, docx_path = self._save_paths(base)
            save_txt(content, txt_path)
            build_styled_doc(
                content=content,
                output_path=docx_path,
                header_image_path=self.header_image_path.get() or None
            )
            self.status.set(f"DOCX 생성 완료 → {docx_path}")
            messagebox.showinfo(APP_NAME, f"문안(TXT/DOCX) 생성 완료\n\n{docx_path}")
        except Exception as e:
            self.status.set("오류")
            messagebox.showerror(APP_NAME, str(e))

    def _pick_image(self):
        path = filedialog.askopenfilename(
            title="배너 이미지 선택",
            filetypes=[("Images","*.png;*.jpg;*.jpeg;*.gif;*.bmp"), ("All files","*.*")]
        )
        if path:
            self.header_image_path.set(path)

# -------------------- 메인 --------------------
if __name__ == "__main__":
    App().mainloop()
